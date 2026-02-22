"""
Generate a fictitious 3-4 page Assessment Report (.docx) for the SuperNova migration (On-prem to Azure).
Structure follows the TOC: Overview, Application Profile, Methodology, Architecture, Infrastructure,
Security and Compliance, Data Management, TCO Assessment & Management.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def main():
    out_dir = Path(__file__).resolve().parent.parent / "sample_docs"
    out_dir.mkdir(exist_ok=True)
    path = out_dir / "SuperNova_Assessment_Report.docx"

    doc = Document()
    doc.add_paragraph()
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("SuperNova Migration Assessment Report")
    r.bold = True
    r.font.size = Pt(18)
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("On-Premises to Microsoft Azure — Technical Assessment").italic = True
    doc.add_paragraph()
    doc.add_paragraph("Confidential — For Internal Use Only").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 1 Overview
    add_heading(doc, "1 Overview", level=1)
    add_para(doc,
        "This document presents the technical assessment for the SuperNova migration initiative, which aims to "
        "move the SuperNova application estate from on-premises infrastructure to Microsoft Azure. SuperNova "
        "currently supports critical order-to-cash and inventory workflows. The assessment covers current state "
        "architecture, target Azure design, infrastructure, security, data management, and total cost of ownership (TCO). "
        "The recommended approach uses Azure Kubernetes Service (AKS), Azure SQL Database, Azure Blob Storage, "
        "and Azure Active Directory (Entra ID) for identity.")
    add_para(doc,
        "The target timeline for Phase 1 (pilot) is Q2 2025, with full production cutover planned for Q4 2025.")

    # 2 Application Profile
    add_heading(doc, "2 Application Profile", level=1)
    add_para(doc,
        "SuperNova comprises a set of .NET Core microservices and a legacy monolith component hosted on Windows Server. "
        "The application integrates with SAP for master data and with internal messaging (currently MSMQ on-prem). "
        "User base is approximately 2,500 concurrent users across North America and EMEA. Data residency requirements "
        "dictate primary Azure region as East US 2 with optional West Europe for DR.")
    add_para(doc,
        "Key dependencies include Active Directory, SQL Server, file shares, and several internal REST APIs. "
        "These are in scope for the migration and will be rehosted or refactored on Azure.")

    # 3 Methodology
    add_heading(doc, "3 Methodology", level=1)
    add_para(doc,
        "The assessment followed a standard discovery-and-design approach: (1) inventory and dependency mapping via "
        "Azure Migrate and discovery tools; (2) application and data classification; (3) target architecture design "
        "using Azure Well-Architected Framework; (4) security and compliance review; (5) TCO and runbook planning. "
        "Stakeholder workshops were held with development, operations, and security teams to validate assumptions "
        "and priorities.")

    # 4 Architecture and Platform (Current & Future)
    add_heading(doc, "4 Architecture and Platform (Current & Future)", level=1)
    add_para(doc,
        "Current state: Three-tier deployment on VMware (web tier, app tier, database tier) with SQL Server Always On "
        "and shared SAN storage. Future state: Web and app tiers containerized and deployed on AKS; Azure SQL Database "
        "for transactional data; Azure Blob Storage for documents and archives; Azure Service Bus for messaging; "
        "Azure Front Door for global load balancing and WAF.")
    add_heading(doc, "4.1 Findings", level=2)
    add_para(doc,
        "The existing .NET Core services are largely cloud-ready with minimal code changes. The legacy monolith "
        "requires either lift-and-shift to Azure VMs initially or a phased refactor into containers. Network topology "
        "allows direct ExpressRoute connectivity from on-prem to Azure; hybrid connectivity will be retained during "
        "transition.")
    add_heading(doc, "4.2 Recommendations", level=2)
    add_para(doc,
        "Adopt AKS for all new and containerized workloads; use Azure Arc for any remaining VM-based components "
        "during transition. Standardize on Azure Monitor, Log Analytics, and Application Insights for observability. "
        "Implement Azure DevOps pipelines for CI/CD targeting AKS and Azure SQL.")

    # 5 Infrastructure
    add_heading(doc, "5 Infrastructure", level=1)
    add_para(doc,
        "Target infrastructure includes AKS clusters (dev, staging, prod) with node pools sized per environment; "
        "Azure SQL Database (Business Critical tier for prod); Azure Blob Storage (Hot and Cool tiers); Virtual Network "
        "with subnets for AKS, data, and management; and Azure Firewall or NVA for egress control.")
    add_heading(doc, "5.1 Findings", level=2)
    add_para(doc,
        "Current on-prem capacity is undersized for peak; Azure auto-scaling will address this. Storage performance "
        "requirements are met by Premium SSD and Azure Blob tiering. ExpressRoute is already in place and can be "
        "extended for Azure connectivity.")
    add_heading(doc, "5.2 Recommendations", level=2)
    add_para(doc,
        "Use Azure Resource Manager templates or Terraform for repeatable deployment. Enable Azure Backup for AKS "
        "persistent volumes and Azure SQL. Define clear tagging and resource-group strategy for cost and governance.")

    # 6 Security and Compliance
    add_heading(doc, "6 Security and Compliance", level=1)
    add_para(doc,
        "Security posture will leverage Azure Entra ID for authentication and Azure RBAC for authorization. Secrets "
        "will be stored in Azure Key Vault. Data at rest and in transit will be encrypted; Azure Defender for "
        "Cloud will be enabled. Compliance scope includes SOC 2 and internal policies.")
    add_heading(doc, "6.1 Findings", level=2)
    add_para(doc,
        "Current on-prem controls are partially aligned with cloud equivalents. Some custom integrations will need "
        "to be reimplemented using Azure-native services (e.g., Key Vault, Managed Identity). No critical compliance "
        "gaps were identified for the target design.")
    add_heading(doc, "6.2 Recommendations", level=2)
    add_para(doc,
        "Implement managed identities for AKS and app services to avoid credential storage. Enable Microsoft Defender "
        "for SQL and Defender for Containers. Conduct a formal security review before production go-live and document "
        "compliance evidence in Azure Policy.")

    # 7 Data Management
    add_heading(doc, "7 Data Management", level=1)
    add_para(doc,
        "SuperNova has an Operational Data Store (ODS) and a Data Warehouse component. The ODS will be migrated to "
        "Azure SQL Database; the warehouse is a candidate for Azure Synapse Analytics or Azure SQL Database with "
        "partitioning for reporting workloads. Azure Data Factory will be used for ETL and scheduled data movement.")
    add_para(doc,
        "Findings and Risks: Data volume is approximately 8 TB for the ODS and 15 TB for the warehouse. Migration "
        "windows and cutover sequencing must be planned to minimize downtime. Data quality and lineage tools should "
        "be extended to Azure.")
    add_heading(doc, "7.1 Recommendations", level=2)
    add_para(doc,
        "Use Azure Database Migration Service or native backup/restore for SQL migration. Implement Azure Purview "
        "for catalog and governance. Retain Azure Blob lifecycle policies for archival and cost optimization.")

    # 8 TCO Assessment & Management
    add_heading(doc, "8 TCO Assessment & Management", level=1)
    add_para(doc,
        "A three-year TCO comparison (on-prem vs. Azure) was performed. Azure costs include compute (AKS, VMs where "
        "needed), database, storage, networking, and licensing (e.g., Azure Hybrid Benefit). Management and tooling "
        "costs are included.")
    add_heading(doc, "8.1 Findings and Risks", level=2)
    add_para(doc,
        "Year 1 shows a modest increase due to dual-run and migration effort; Year 2 and 3 show an estimated 15–20% "
        "reduction in run cost with improved scalability and reduced hardware refresh. Key risks include unplanned "
        "data egress and over-provisioned resources; cost management and budgets will be enforced via Azure Cost "
        "Management and tags.")
    add_heading(doc, "8.2 Recommendations", level=2)
    add_para(doc,
        "Establish monthly cost reviews and set budget alerts. Use Azure Reservations and Savings Plans for committed "
        "workloads. Assign a dedicated FinOps role to track and optimize SuperNova-related spend post-migration.")

    doc.add_paragraph()
    doc.add_paragraph("— End of Assessment Report —").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(path)
    print(f"Created: {path}")
    return path

if __name__ == "__main__":
    main()
