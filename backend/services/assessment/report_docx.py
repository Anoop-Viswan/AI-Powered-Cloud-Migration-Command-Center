"""Convert assessment report (markdown or plain text) to DOCX for download."""

import re
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _diagram_png_path(assessment_id: str) -> Path | None:
    """Path to generated target architecture PNG, if it exists."""
    root = Path(__file__).resolve().parent.parent.parent.parent
    p = root / "data" / "assessment_diagrams" / assessment_id / "target_architecture.png"
    return p if p.exists() else None


def report_to_docx(
    report_text: str,
    title: str = "Migration Assessment Report",
    assessment_id: str | None = None,
) -> bytes:
    """
    Convert report text (markdown or plain) to DOCX bytes.
    If assessment_id is provided and the report contains the target diagram image,
    the diagram PNG is embedded in the DOCX so the downloaded report shows the diagram.
    """
    doc = Document()
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(18)
    doc.add_paragraph()
    doc.add_paragraph("Confidential — For Internal Use Only").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    if not (report_text or "").strip():
        doc.add_paragraph("(No report content.)")
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # Split into blocks: treat ## or # as heading, else paragraph
    lines = report_text.replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        # Markdown heading: # or ## or ###
        if stripped.startswith("#"):
            level = 0
            while level < len(stripped) and stripped[level] == "#":
                level += 1
            heading_text = stripped[level:].strip()
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(14 if level == 1 else 12)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        # Image markdown: ![alt](url) – embed diagram PNG if we have assessment_id and the file exists
        if re.match(r"^!\[.*\]\(.*\)\s*$", stripped):
            if assessment_id:
                diagram_path = _diagram_png_path(assessment_id)
                if diagram_path:
                    try:
                        doc.add_picture(str(diagram_path), width=Inches(5.5))
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(6)
                    except Exception:
                        p = doc.add_paragraph("[Target State Architecture diagram – view in app]")
                        p.paragraph_format.space_after = Pt(6)
                else:
                    p = doc.add_paragraph("[Target State Architecture diagram – generate report to export]")
                    p.paragraph_format.space_after = Pt(6)
            else:
                p = doc.add_paragraph("[Diagram – see report in app]")
                p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        # Collect paragraph (until blank or next # or image line)
        para_lines = []
        while i < len(lines):
            ln = lines[i]
            if ln.strip().startswith("#") or (ln.strip() == "" and para_lines):
                break
            if re.match(r"^!\[.*\]\(.*\)\s*$", ln.strip()):
                break
            if ln.strip():
                para_lines.append(ln.strip())
            i += 1
        if para_lines:
            text = " ".join(para_lines)
            # Strip markdown bold/italic for plain DOCX
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
